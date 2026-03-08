"""Tests for create_docx tool."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from office_agent.tools.docx_tool import create_docx


def test_create_docx_basic(sample_docx_plan: dict, tmp_path: Path) -> None:
    result = create_docx(plan=sample_docx_plan, out_dir=str(tmp_path))
    assert result["tool_used"] == "create_docx"
    out = Path(result["output_path"])
    assert out.exists()
    assert out.suffix == ".docx"


def test_create_docx_content(sample_docx_plan: dict, tmp_path: Path) -> None:
    result = create_docx(plan=sample_docx_plan, out_dir=str(tmp_path))
    doc = Document(result["output_path"])
    paragraphs_text = [p.text for p in doc.paragraphs]
    full_text = "\n".join(paragraphs_text)
    assert sample_docx_plan["title"] in full_text


def test_create_docx_has_table(sample_docx_plan: dict, tmp_path: Path) -> None:
    result = create_docx(plan=sample_docx_plan, out_dir=str(tmp_path))
    doc = Document(result["output_path"])
    assert len(doc.tables) >= 1


def test_create_docx_missing_title(tmp_path: Path) -> None:
    plan = {"sections": [{"heading": "テスト", "level": 1}]}
    result = create_docx(plan=plan, out_dir=str(tmp_path))
    out = Path(result["output_path"])
    assert out.exists()


def test_create_docx_output_dir_created(sample_docx_plan: dict, tmp_path: Path) -> None:
    nested = tmp_path / "a" / "b" / "c"
    result = create_docx(plan=sample_docx_plan, out_dir=str(nested))
    assert Path(result["output_path"]).exists()


def test_tool_registry_docx(registry, sample_docx_plan: dict, tmp_path: Path) -> None:
    result = registry.call("create_docx", plan=sample_docx_plan, out_dir=str(tmp_path))
    assert result["tool_used"] == "create_docx"
