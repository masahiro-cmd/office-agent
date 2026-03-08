"""create_docx: Generate a Word document from a plan JSON."""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.shared import Pt

logger = logging.getLogger(__name__)

# Heading level → Word style name mapping
_HEADING_STYLES = {
    1: "Heading 1",
    2: "Heading 2",
    3: "Heading 3",
}


def create_docx(plan: dict, out_dir: str, template_dir: str = "./templates") -> dict:
    """
    Generate a .docx file from a validated plan dict.

    Args:
        plan: Validated document plan (must conform to docx_schema.json).
        out_dir: Directory where the output file will be saved.
        template_dir: Directory containing optional .docx template files.

    Returns:
        {"output_path": str, "tool_used": "create_docx"}
    """
    title: str = plan.get("title", "Untitled")
    sections: list[dict] = plan.get("sections", [])

    # Check for a template file
    template_path = _find_template(template_dir)
    if template_path:
        logger.info(f"Using template: {template_path}")
        doc = Document(str(template_path))
    else:
        doc = Document()

    # Title
    doc.add_heading(title, level=0)

    # Metadata paragraph
    meta: dict = plan.get("metadata", {})
    if meta:
        meta_text = "  |  ".join(
            f"{k}: {v}" for k, v in meta.items() if v
        )
        meta_para = doc.add_paragraph(meta_text)
        meta_para.style = "Normal"
        for run in meta_para.runs:
            run.font.size = Pt(9)

    # Sections
    for section in sections:
        _add_section(doc, section)

    # Save
    out_path = Path(out_dir)
    out_path.mkdir(parents=True, exist_ok=True)
    safe_title = _safe_filename(title)
    file_path = out_path / f"{safe_title}.docx"
    doc.save(str(file_path))
    logger.info(f"Saved docx: {file_path}")

    return {"output_path": str(file_path), "tool_used": "create_docx"}


def _add_section(doc: Document, section: dict) -> None:
    heading = section.get("heading", "")
    level = int(section.get("level", 1))
    content = section.get("content", "")
    bullets = section.get("bullets") or []
    table_data = section.get("table")
    footnote = section.get("footnote")

    if heading:
        doc.add_heading(heading, level=level)

    if content:
        doc.add_paragraph(content)

    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")

    if table_data:
        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])
        if headers:
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            # Header row
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = str(h)
            # Data rows
            for r_idx, row in enumerate(rows):
                cells = table.rows[r_idx + 1].cells
                for c_idx, val in enumerate(row):
                    cells[c_idx].text = str(val)

    if footnote:
        fn_para = doc.add_paragraph(footnote)
        for run in fn_para.runs:
            run.font.size = Pt(9)


def _find_template(template_dir: str) -> Path | None:
    td = Path(template_dir) / "word"
    if td.exists():
        for ext in ("*.docx",):
            matches = list(td.glob(ext))
            if matches:
                return matches[0]
    return None


def _safe_filename(name: str) -> str:
    """Replace characters that are invalid in filenames."""
    invalid = r'\/:*?"<>|'
    for ch in invalid:
        name = name.replace(ch, "_")
    return name[:100]
